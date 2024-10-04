from io import BytesIO
from pathlib import Path
from typing import Set, Union

import docx
from docling_core.types.experimental import (
    BasePictureData,
    BaseTableData,
    DescriptionItem,
    DocItemLabel,
    DoclingDocument,
    GroupLabel,
    ImageRef,
    PictureItem,
    SectionHeaderItem,
    TableCell,
    TableItem,
)
from lxml import etree

from docling.backend.abstract_backend import DeclarativeDocumentBackend
from docling.datamodel.base_models import InputFormat


class MsWordDocumentBackend(DeclarativeDocumentBackend):

    def __init__(self, path_or_stream: Union[BytesIO, Path], document_hash: str):
        self.XML_KEY = (
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val"
        )
        super().__init__(path_or_stream, document_hash)
        # self.initialise(path_or_stream)
        # Word file:
        self.path_or_stream = path_or_stream
        # Initialise the parents for the hierarchy
        self.max_levels = 10
        self.level_at_new_list = None
        self.parents = {}
        for i in range(-1, self.max_levels):
            self.parents[i] = None

        self.history = {
            "names": [None],
            "levels": [None],
            "numids": [None],
            "indents": [None],
        }

    def warn(self, message):
        print(f"WARN: {message}")

    def is_valid(self) -> bool:
        return True

    def is_paginated(cls) -> bool:
        False

    def unload(self):
        if isinstance(self.path_or_stream, BytesIO):
            self.path_or_stream.close()

        self.path_or_stream = None

    def update_history(self, name, level, numid, ilevel):
        self.history["names"].append(name)
        self.history["levels"].append(level)

        self.history["numids"].append(numid)
        self.history["indents"].append(ilevel)

    def prev_name(self):
        return self.history["names"][-1]

    def prev_level(self):
        return self.history["levels"][-1]

    def prev_numid(self):
        return self.history["numids"][-1]

    def prev_indent(self):
        return self.history["indents"][-1]

    def get_level(self) -> int:
        """Return the first None index."""
        for k, v in self.parents.items():
            if k >= 0 and v == None:
                return k
        return 0

    @classmethod
    def supported_formats(cls) -> Set[InputFormat]:
        return {InputFormat.DOCX}

    def walk_linear(self, body, docx_obj, doc) -> DoclingDocument:
        for element in body:
            tag_name = etree.QName(element).localname
            # Check for Text
            if tag_name in ["p"]:
                self.add_text(element, docx_obj, doc)
            # Check for Tables
            elif element.tag.endswith("tbl"):
                self.add_table(element, docx_obj, doc)
            # Check for Inline Images (drawings or blip elements)
            elif element.tag.endswith("drawing") or element.tag.endswith("blip"):
                self.add_figure(element, docx_obj, doc)
            else:
                self.warn(f"ignoring element in DOCX with tag: {tag_name}")
        return doc

    def get_numId_and_ilvl(self, paragraph):
        # Access the XML element of the paragraph
        numPr = paragraph._element.find(
            ".//w:numPr", namespaces=paragraph._element.nsmap
        )

        if numPr is not None:
            # Get the numId element and extract the value
            numId_elem = numPr.find("w:numId", namespaces=paragraph._element.nsmap)
            ilvl_elem = numPr.find("w:ilvl", namespaces=paragraph._element.nsmap)

            numId = numId_elem.get(self.XML_KEY) if numId_elem is not None else None
            ilvl = ilvl_elem.get(self.XML_KEY) if ilvl_elem is not None else None

            return int(numId), int(ilvl)

        return None, None  # If the paragraph is not part of a list

    def get_label_and_level(self, paragraph):
        if paragraph.style is None:
            return "Normal", None
        label = paragraph.style.name
        if ":" in label:
            parts = label.split(":")

            if len(parts) == 2:
                return parts[0], int(parts[1])

        parts = label.split(" ")

        if "Heading" in label and len(parts) == 2:
            return parts[0], int(parts[1])
        else:
            return label, None

    def add_text(self, element, docx_obj, doc):
        paragraph = docx.text.paragraph.Paragraph(element, docx_obj)

        if paragraph.text is None:
            # self.warn(f"paragraph has text==None")
            return

        text = paragraph.text.strip()

        # if len(text)==0: # keep empty paragraphs, they seperate adjacent lists!
        #     self.warn(f"paragraph has len(text)==0")

        pname, plevel = self.get_label_and_level(paragraph)
        numid, ilevel = self.get_numId_and_ilvl(paragraph)

        # we detected a list
        if numid is not None and ilevel is not None:
            self.add_listitem(
                element, docx_obj, doc, pname, plevel, numid, ilevel, text
            )
            self.update_history(pname, plevel, numid, ilevel)

            return

        elif numid is None and self.prev_numid() is not None:  # Close list

            for key, val in self.parents.items():
                if key >= self.level_at_new_list:
                    self.parents[key] = None

            self.level = self.level_at_new_list - 1
            self.level_at_new_list = None

        if pname in ["Title"]:

            for key, val in self.parents.items():
                self.parents[key] = None

            self.parents[0] = doc.add_text(
                parent=None, label=DocItemLabel.TITLE, text=text
            )

        elif "Heading" in pname:
            self.add_header(element, docx_obj, doc, pname, plevel, text)

        elif pname in [
            "Paragraph",
            "Normal",
            "Subtitle",
            "Author",
            "Default Text",
            "List Paragraph",
            "List Bullet",
            "Quote",
        ]:
            level = self.get_level()
            doc.add_text(
                label=DocItemLabel.PARAGRAPH, parent=self.parents[level - 1], text=text
            )

        else:
            assert False, f"need to add a new paragraph: {pname}"

        self.update_history(pname, plevel, numid, ilevel)

    def add_header(self, element, docx_obj, doc, curr_name, curr_level, text: str):

        level = self.get_level()
        # print(f"level: {level} => add_header(self, element, docx_obj, doc, {curr_name}, {curr_level}): {text}")

        if isinstance(curr_level, int):

            if curr_level == level:

                self.parents[level] = doc.add_heading(
                    parent=self.parents[level - 1], text=text
                )

            elif curr_level > level:

                # add invisible group
                for i in range(level, curr_level):
                    self.parents[i] = doc.add_group(
                        parent=self.parents[i - 1],
                        label=GroupLabel.SECTION,
                        name=f"header-{i}",
                    )

                self.parents[curr_level] = doc.add_heading(
                    parent=self.parents[curr_level - 1], text=text
                )

            elif curr_level < level:

                # remove the tail
                for key, val in self.parents.items():
                    if key >= curr_level:
                        self.parents[key] = None

                self.parents[curr_level] = doc.add_heading(
                    parent=self.parents[curr_level - 1], text=text
                )

        else:
            self.parents[self.level] = doc.add_heading(
                parent=self.parents[self.level - 1], text=text
            )

    def add_listitem(
        self, element, docx_obj, doc, pname, plevel, numid, ilevel, text: str
    ):

        level = self.get_level()
        # print(f"level: {level} => add_listitem(self, element, docx_obj, doc, {pname}, {plevel}, {numid}, {ilevel}): {text}")

        if self.prev_numid() is None:  # Open new list

            self.level_at_new_list = level

            self.parents[level] = doc.add_group(
                label=GroupLabel.LIST, name="list", parent=self.parents[level - 1]
            )

            doc.add_text(
                label=DocItemLabel.LIST_ITEM, parent=self.parents[level], text=text
            )

        elif (
            self.prev_numid() == numid and self.prev_indent() < ilevel
        ):  # Open indented list

            for i in range(
                self.level_at_new_list + self.prev_indent() + 1,
                self.level_at_new_list + ilevel + 1,
            ):
                self.parents[i] = doc.add_group(
                    label=GroupLabel.LIST, name="list", parent=self.parents[i - 1]
                )

            doc.add_text(
                label=DocItemLabel.LIST_ITEM,
                parent=self.parents[self.level_at_new_list + ilevel],
                text=text,
            )

        elif self.prev_numid() == numid and ilevel < self.prev_indent():  # Close list

            for k, v in self.parents.items():
                if k > self.level_at_new_list + ilevel:
                    self.parents[k] = None

            doc.add_text(
                label=DocItemLabel.LIST_ITEM,
                parent=self.parents[self.level_at_new_list + ilevel],
                text=text,
            )

        elif self.prev_numid() == numid or self.prev_indent() == ilevel:

            doc.add_text(
                label=DocItemLabel.LIST_ITEM, parent=self.parents[level - 1], text=text
            )

    def add_table(self, element, docx_obj, doc):

        # Function to check if a cell has a colspan (gridSpan)
        def get_colspan(cell):
            grid_span = cell._element.xpath("@w:gridSpan")
            if grid_span:
                return int(grid_span[0])  # Return the number of columns spanned
            return 1  # Default is 1 (no colspan)

        # Function to check if a cell has a rowspan (vMerge)
        def get_rowspan(cell):
            v_merge = cell._element.xpath("@w:vMerge")
            if v_merge:
                return v_merge[
                    0
                ]  # 'restart' indicates the beginning of a rowspan, others are continuation
            return 1

        table = docx.table.Table(element, docx_obj)

        num_rows = len(table.rows)
        num_cols = 0
        for row in table.rows:
            # Calculate the max number of columns
            num_cols = max(num_cols, sum(get_colspan(cell) for cell in row.cells))

        self.warn(f"table: [{num_rows}x{num_cols}]")

        # Initialize the table grid
        table_grid = [[None for _ in range(num_cols)] for _ in range(num_rows)]

        data = BaseTableData(num_rows=num_rows, num_cols=num_cols, table_cells=[])

        for row_idx, row in enumerate(table.rows):
            col_idx = 0
            for c, cell in enumerate(row.cells):
                row_span = get_rowspan(cell)
                col_span = get_colspan(cell)

                # Find the next available column in the grid
                while table_grid[row_idx][col_idx] is not None:
                    col_idx += 1

                print(f"{row_idx}, {col_idx}, {row_span}, {col_span}")

                # Fill the grid with the cell value, considering rowspan and colspan
                for i in range(row_span if row_span == "restart" else 1):
                    for j in range(col_span):
                        table_grid[row_idx + i][col_idx + j] = ""

                cell = TableCell(
                    text=cell.text,
                    row_span=row_span,
                    col_span=col_span,
                    start_row_offset_idx=row_idx,
                    end_row_offset_idx=row_idx + row_span,
                    start_col_offset_idx=col_idx,
                    end_col_offset_idx=col_idx + col_span,
                    col_header=False,  # col_header,
                    row_header=False,  # ((not col_header) and html_cell.name=='th')
                )

                # print(row_idx, "\t", col_idx, "\t", row_span, "\t", col_span, "\t", text)
                data.table_cells.append(cell)

        level = self.get_level()
        doc.add_table(data=data, parent=self.parents[level - 1])

    def add_figure(self, element, docx_obj, doc):
        doc.add_picture(
            data=BasePictureData(), parent=self.parents[self.level], caption=None
        )

    def convert(self) -> DoclingDocument:
        # Parses the DOCX into a structured document model.
        doc = DoclingDocument(description=DescriptionItem(), name="dummy")
        docx_obj = None
        try:
            docx_obj = docx.Document(self.path_or_stream)
        except Exception:
            return doc

        # self.initialise()
        doc = self.walk_linear(docx_obj.element.body, docx_obj, doc)

        return doc